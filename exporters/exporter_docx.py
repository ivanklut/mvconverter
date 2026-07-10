# Концептуальная схема будущего DOCX-экспортера Марквана
class DocxExporter:
	def __init__(self):
		# Твое железное правило: Зеркальный маппинг классов и функций! [INDEX]
		self.render_map = {
			'Section': self._Section,
			'Heading': self._Heading,
			'Paragraph': self._Paragraph,
			'EndSection': self._EndSection
		}

	def export_to_docx(self, doc: Document) -> bytes:
		"""Главный метод. Создает документ Word в памяти и возвращает его байты [INDEX]"""
		import io
		from docx import Document as WordDocument
		
		word_doc = WordDocument()
		
		# Добавляем заголовок книги на основе метаданных UNO [INDEX]
		book_title = doc.metadata.get("title", "Документ Маркван")
		word_doc.add_heading(book_title, level=0)
		
		# Запускаем конвейер обхода хронологического потока нод [INDEX]
		self._render_nodes_flow(doc.body.nodes, word_doc)
		
		# Сейвим документ не на жесткий диск, а в виртуальный буфер памяти байт! [INDEX]
		docx_buffer = io.BytesIO()
		word_doc.save(docx_buffer)
		
		return docx_buffer.getvalue() # Отдаем чистые бинарные байты файла .docx Билдеру! [INDEX]

	def _render_nodes_flow(self, nodes_list, word_doc):
		"""Линейно бежит по нодам и вызывает их зеркальные методы из таблицы [INDEX]"""
		for n in nodes_list:
			class_name = n.__class__.__name__
			handler = self.render_map.get(class_name)
			if handler:
				handler(n, word_doc)

	# --- КОМПАКТНЫЕ МЕТОДЫ-ОБРАБОТЧИКИ БЛОКОВ ---

	def _Section(self, n, word_doc):
		"""Рекурсивно уходит вглубь секций-матрешек [INDEX]"""
		self._render_nodes_flow(n.nodes, word_doc)

	def _Heading(self, n, word_doc):
		"""Добавляет заголовок нужного уровня в Word (пересчитывая n.level) [INDEX]"""
		word_doc.add_heading(n.text, level=n.level if n.level <= 9 else 9)

	def _Paragraph(self, n, word_doc):
		"""Создает абзац и наполняет его богатым текстом из n.inlines! [INDEX]"""
		p = word_doc.add_paragraph()
		# Вызываем быстрый обходчик внутристрочных стилей
		self._render_inline_flow(n.inlines, p)

	def _EndSection(self, n, word_doc):
		"""Маркер ___ превращается в каноничный разрыв страницы в Word! [INDEX]"""
		word_doc.add_page_break()
